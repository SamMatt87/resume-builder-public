import tkinter
from tkinter import ttk
import datetime
import extractions
from docx import Document
import os
import time
class ResumeForm:
	def __init__(self, master, skill_set,tech, skills_list, role_summaries, pu_role_summaries, companies, tasks, rtags, degrees, courses, tags, degree_summaries, grades):
		self.label = tkinter.Label(master, text='Resume Builder').pack()
		self.Company = ttk.Entry(master,width=50)
		self.Company.pack()
		self.contacts = tkinter.Frame(master)
		self.inc_email = tkinter.BooleanVar(value = True)
		self.email = ttk.Checkbutton(self.contacts, text='Email', onvalue = 1, offvalue = 0, variable = self.inc_email)
		self.email.pack(side='left')
		self.inc_phone = tkinter.BooleanVar()
		self.phone = ttk.Checkbutton(self.contacts, text='Phone',onvalue = 1, offvalue = 0, variable = self.inc_phone)
		self.phone.pack(side='left')
		self.inc_git = tkinter.BooleanVar(value = True)
		self.git = ttk.Checkbutton(self.contacts, text='Github',onvalue = 1, offvalue = 0, variable = self.inc_git)
		self.git.pack(side='left')
		self.inc_web = tkinter.BooleanVar()
		self.website = ttk.Checkbutton(self.contacts, text='Website',onvalue = 1, offvalue = 0, variable = self.inc_web)
		self.website.pack(side='left')
		self.contacts.pack()
		self.skills_roles = tkinter.Frame(master)
		self.skills = tkinter.Frame(self.skills_roles)
		tkinter.Label(self.skills, text='Skills').pack()
		self.skillbox = tkinter.Listbox(self.skills, selectmode='multiple', exportselection=0)
		for num, skill in zip(range(0,len(skill_set)), sorted(skill_set, key = lambda s : s.lower())):
			self.skillbox.insert(num, skill)
		self.skill_scroll = tkinter.Scrollbar(self.skills, orient='vertical', command=self.skillbox.yview)
		self.skillbox['yscrollcommand'] = self.skill_scroll.set
		self.skillbox.pack(side='left')
		self.skill_scroll.pack(side='right', fill='y')
		self.skills.pack(side='left')
		self.roles = tkinter.Frame(self.skills_roles)
		tkinter.Label(self.roles, text='Roles').pack()
		self.pre_roles = tkinter.BooleanVar()
		self.preuni = tkinter.Checkbutton(self.roles, text="Include pre university roles",onvalue=1, offvalue=0,variable = self.pre_roles,
			command=lambda: self.Additional_roles(role_summaries,pu_role_summaries))
		self.preuni.pack()
		self.rolebox = tkinter.Listbox(self.roles, selectmode='multiple', exportselection=0)
		for num, role in zip(range(0, len(role_summaries)),role_summaries):
			self.rolebox.insert(num,role[0])
		self.rolebox.pack()
		self.roles.pack(side='right')
		self.skills_roles.pack()
		self.degreesbox = tkinter.Listbox(master, selectmode='multiple', exportselection=0)
		for num, degree in zip(range(0,len(set(degrees))),set(degrees)):
			self.degreesbox.insert(num, degree)
		self.degreesbox.pack()
		today = datetime.date.today()
		self.monthyear = tkinter.Frame(master)
		months='January February March April May June July August September October November December'.split()
		self.monthname = ttk.Combobox(self.monthyear, values = months)
		self.monthname.current(today.month-1)
		self.year_var = tkinter.StringVar(self.monthyear)
		self.year_var.set(today.year)
		self.year = ttk.Spinbox(self.monthyear, from_=2020, to = 2060, textvariable = self.year_var)
		self.monthname.pack(side='left')
		self.year.pack(side='right')
		self.monthyear.pack()
		self.button = ttk.Button(master, text='Generate Resume', 
			command=lambda: self.Generate_Resume(skill_set,tech,skills_list,role_summaries, pu_role_summaries, degree_summaries, companies, tasks, rtags,
				degrees, tags, courses, grades))
		self.button.pack()
		self.progress = ttk.Progressbar(master, orient = 'horizontal', length = 300)
		self.progress.pack()
		today = datetime.date.today()
	def Generate_Resume(self,skill_set,tech,skills_list, role_summaries, pu_role_summaries,degree_summaries, companies, tasks, rtags, degrees, tags, courses, grades):
		self.button.config(text='Generating Resume')
		self.button.state(['disabled'])
		self.progress.config(mode='determinate')
		self.Company_name = self.Company.get()
		resume_name = 'Sam Matthews Resume '+' '.join([self.Company.get(), self.monthname.get(), self.year.get()])+'.docx'
		print('Generating '+resume_name)
		doc = Document()
		self.progress.step(14)
		doc.add_heading('Sam Matthews',0)
		contacts =[]
		if self.inc_email.get() ==1:
			contacts.append('Email: Sam.Matthews.77@gmail.com')
		if self.inc_phone.get() ==1:
			contacts.append('Phone: 0423741328')
		if self.inc_git.get() ==1:
			contacts.append('GitHub: https://github.com/SamMatt87')
		if self.inc_web.get() ==1:
			contacts.append('Website: https://sammatt87.github.io/')
		print(' | '.join(contacts))
		doc.add_paragraph(' | '.join(contacts))
		self.progress.step(14)
		technology_set = set()
		skill_selections= [self.skillbox.get(i) for i in self.skillbox.curselection()]
		for skill_num in range(0,len(skills_list)):
			for skill in skills_list[skill_num]:
				if skill in skill_selections:
					technology_set.update([tech[skill_num]])
		print('tecnical skills')
		doc.add_heading('Technical Skills',level=1)
		doc.add_paragraph(' | '.join(technology_set))
		for technology in technology_set:
			print(technology)
		self.progress.step(14)
		role_selections = [self.rolebox.get(i) for i in self.rolebox.curselection()]
		if len(role_selections) >=1:
			print('\nrelevant roles')
			doc.add_heading('Relevant Roles', level=1)
			all_roles = role_summaries+pu_role_summaries
			for role in all_roles:
				if role[0] in role_selections:
					print(role)
					role_tasks = []
					doc.add_heading(role[0]+' - '+role[1]+' - '+role[2]+' to '+role[3],level=2)
					for task_num in range(0,len(tasks)):
						if companies[task_num]==role[0]:
							for rtag in rtags[task_num]:
								if rtag in skill_selections and tasks[task_num] not in role_tasks:
									doc.add_paragraph(tasks[task_num], style='List Bullet')
									role_tasks.append(tasks[task_num])
			if self.rolebox.size() > len(role_selections):
				print('\nother roles')
				doc.add_heading('Other Roles', level=1)
		else:
			print('\nprevious roles')
			doc.add_heading('Roles', level=1)
		self.progress.step(14)
		for role in role_summaries:
			if role[0] not in role_selections:
				print(role)
				doc.add_paragraph(role[0]+' - '+role[1]+' - '+role[2]+' to '+role[3])
		if self.pre_roles.get()==1:
			for role in pu_role_summaries:
				if role[0] not in role_selections:
					print(role)
					doc.add_paragraph(role[0]+' - '+role[1]+' - '+role[2]+' to '+role[3])
		self.progress.step(14)
		degree_selections = [self.degreesbox.get(i) for i in self.degreesbox.curselection()]
		doc.add_heading('Education', level=1)
		for degree in degree_summaries:
			print(degree[0]+' - '+degree[2]+' - '+degree[3]+' to '+degree[4])
			doc.add_heading(degree[0]+' - '+degree[2]+' - '+degree[3]+' to '+degree[4])
			if degree[0] in degree_selections:
				subject_list = []
				print('relevant subjects')
				doc.add_heading('Relevant Subjects',level=2)
				for sub_num in range(0,len(courses)):
					if degrees[sub_num]==degree[0]:
						for tag in tags[sub_num]:
							if tag in skill_selections and courses[sub_num] not in subject_list:
								print(courses[sub_num]+' - '+str(grades[sub_num]))
								doc.add_paragraph(courses[sub_num]+' - '+str(grades[sub_num]), style='List Bullet')
								subject_list.append[courses[sub_num]]
		self.progress.step(14)
		doc.save(resume_name)
		self.progress.step(15)
		time.sleep(2)
		os.startfile(resume_name,'open')
	def Additional_roles(self,role_summaries, pu_role_summaries):
		if self.pre_roles.get()==1:
			for role_num, role in zip(range(0,len(pu_role_summaries)),pu_role_summaries):
				self.rolebox.insert(len(role_summaries)+role_num,role[0])
		elif self.pre_roles.get()==0:
			for role_num in range(0,len(pu_role_summaries)):
				self.rolebox.delete(len(role_summaries))

def main():
	print('start')
	now = datetime.datetime.now()
	degree_summaries, degrees, courses, grades, tags, unique_tags= extractions.extract_courses()
	print(datetime.datetime.now()-now)
	tech, skills, skill_set = extractions.extract_skills()
	print(datetime.datetime.now()-now)
	role_summaries, pu_role_summaries, companies, tasks, rtags, unique_rtags= extractions.exract_roles()
	print(datetime.datetime.now()-now)
	skill_set.update(unique_rtags,unique_tags)
	print(datetime.datetime.now()-now)
	root = tkinter.Tk()
	app = ResumeForm(root, skill_set, tech, skills, role_summaries, pu_role_summaries, companies, tasks, rtags, degrees,courses, tags, degree_summaries, grades)
	print(datetime.datetime.now()-now)
	root.mainloop()

if __name__ == "__main__":
	main()